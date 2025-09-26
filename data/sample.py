import os
import re
import pandas as pd
from pdf2image import convert_from_path
import pytesseract
from docx import Document
import cv2

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
poppler_path = r"D:\Release-25.07.0-0\poppler-25.07.0\Library\bin"  
def extract_sections(pdf_path):
    text = ""
    try:
        # Convert PDF pages to images
        pages = convert_from_path(pdf_path, dpi=300,poppler_path=poppler_path)
        for page in pages:
           
            page_text = pytesseract.image_to_string(page, lang="eng")
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        print(f"Error processing {pdf_path}: {e}")
        return []

    # Split into sections (headings detection)
    sections = re.split(r"\n(?=[A-Z][A-Za-z\s]{3,50})", text)
    return sections

pdf_files = [
    {"path": r"D:\info\data\india\NEP.pdf", 
     "country": "India", "domain": "Education", "title": "National Education Policy", "year": 2021},
    
]

records = []
for pdf in pdf_files:
    sections = extract_sections(pdf["path"])
    for sec in sections:
        sec = sec.strip()
        heading_candidate = sec.split("\n")[0]
        if heading_candidate.isupper() and len(heading_candidate.split()) < 10:
            heading = heading_candidate[:80]
            content = "\n".join(sec.split("\n")[1:])
            records.append({
                "Country": pdf["country"],
                "Domain": pdf["domain"],
                "Title": pdf["title"],
                "Year": pdf["year"],
                "Section": heading,
                "Content": content
            })

df = pd.DataFrame(records)
df.to_csv("tesseract_policies_dataset3.csv", index=False, encoding="utf-8")
doc = Document()
doc.add_heading("Policies Dataset", 0)
doc.add_page_break()

for rec in records:
    doc.add_page_break()
    doc.add_heading(f"{rec['Title']} ({rec['Country']}, {rec['Year']})", level=1)

    p = doc.add_paragraph()
    p.add_run("Domain: ").bold = True
    p.add_run(rec["Domain"] + "\n")
    p.add_run("Country: ").bold = True
    p.add_run(rec["Country"] + "\n")
    p.add_run("Year: ").bold = True
    p.add_run(str(rec["Year"]))

    doc.add_heading(rec["Section"], level=2)

    # Preserve paragraphs instead of merging
    for para in rec["Content"].split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

output_path = r"d:\info\data\tesseract_policies_dataset3.docx"
print("Saving to:", output_path)
doc.save(output_path)

print("Dataset saved as tesseract_policies_dataset3.csv and tesseract_policies_dataset3.docx with", len(df), "rows")
